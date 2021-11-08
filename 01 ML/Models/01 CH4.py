# Load library
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns; sns.set(style = 'ticks', color_codes = True)

from sklearn.preprocessing import StandardScaler
from sklearn.decomposition import PCA
from sklearn.model_selection import train_test_split, GridSearchCV, KFold
from sklearn.metrics import make_scorer, mean_squared_error, r2_score
from sklearn.linear_model import ElasticNet, Lasso
from sklearn.neighbors import KNeighborsRegressor
from sklearn.svm import SVR
from sklearn.tree import DecisionTreeRegressor
from sklearn.ensemble import RandomForestRegressor
from xgboost import XGBRegressor, plot_importance, plot_tree
from tensorflow.keras.models import Sequential
from tensorflow.keras.layers import Dense
from tensorflow.keras.wrappers.scikit_learn import KerasRegressor
from docx import Document

#ignore warnings during runing models
import warnings
warnings.filterwarnings('ignore')

# Load datasets
df = pd.read_csv('C:\\Users\\Administrator\\Desktop\\20200818\\01 ML\\CH4.csv').drop(['No'], axis = 1)

# Draw distplot of uptake amount
plt.figure(figsize = (8, 8))
plt.xlabel('Methane Uptake', fontsize = 18)
plt.ylabel('Percentage', fontsize = 18)
sns.distplot(df['G'], axlabel = 'Methane Uptake', label = 'Percentage')
plt.savefig('F1 Distplot.png', dpi = 300)
plt.show()

# Draw pairplot and heatmap
plt.figure(figsize = (8, 8))
sns.pairplot(df.drop(['G'], axis = 1), palette = 'husl', plot_kws=dict(edgecolor = None))
plt.savefig('F2 Pairplot.png', dpi = 300)
plt.show()

plt.figure(figsize = (10, 8))
ax = sns.heatmap(df.drop(['G'], axis = 1).corr(), cmap="YlGnBu", annot = True, vmin = 0, vmax = 1)
ax.set_ylim([5, 0])
plt.savefig('F3 Heatmap.png', dpi = 300)
plt.show()

subset = df[(df['T'] == 25) & (df['P'] == 1).reset_index(drop = True)]
subset = subset.iloc[:, [1, 2, 5]]
plt.figure(figsize = (10, 6))
plt.tricontourf(subset.iloc[:, 1].values,
                subset.iloc[:, 0].values,
                subset.iloc[:, 2].values,
                cmap = 'jet',
                levels = 100)
plt.xlabel('Vmeso', fontsize = 18)
plt.ylabel('Vmicro', fontsize = 18)
plt.colorbar()
plt.ylim([0, 1])
plt.savefig('F4 Vmeso vs Vmicro', dpi = 300)
plt.show()

for rs in [7, 14, 21, 35, 42, 84, 100, 120, 200, 420]:
    # Split the datasets into training and test datasets
    X = df.drop(['G'], axis = 1)
    y = df['G']
    
    X_train, X_test, y_train, y_test = train_test_split(X, y, 
                                                        test_size = 0.2, 
                                                        shuffle = True, 
                                                        random_state = rs)
    y_train = np.asarray(y_train)
    y_test = np.asarray(y_test)
    # Standardization
    sds = StandardScaler()
    X_train_scaled = sds.fit_transform(X_train)
    X_test_scaled = sds.transform(X_test)
    # PCA check
    pca = PCA(n_components = 2)
    X_train_pca = pca.fit_transform(X_train_scaled)
    X_test_pca = pca.fit_transform(X_test_scaled)
    ## Visualize PCA 
    plt.figure(figsize = (8, 8))
    plt.scatter(X_train_pca[:, 0], X_train_pca[:, 1], c='Tab:purple', alpha = 0.6, edgecolors = 'none',)
    plt.scatter(X_test_pca[:, 0], X_test_pca[:, 1], c='Tab:blue', alpha = 0.6, edgecolors = 'none',)
    plt.xlabel('Component 1', fontsize = 18)
    plt.ylabel('Component 2', fontsize = 18)
    plt.legend(loc="upper right", frameon = False)
    plt.savefig('F5 PCA Visualization RS{}.png'.format(rs), dpi = 300)
    plt.show()
    
    # Cross-validation
    cv = KFold(n_splits = 5, random_state = rs)

    # Evaluation metrics
    scoring = make_scorer(mean_squared_error, greater_is_better=False)
    # output score of trained model 
    def save_evaluate_draw(y_pred_train, y_pred_test, grid_param, best_params, best_estimator, model):
        # Save groundtruth and prediction results into excel
        train_array = pd.DataFrame({'y_train': np.array(y_train),
                                    'y_pred_train': y_pred_train,
                                    'AE':np.abs(y_train - y_pred_train),
                                    'APE': np.abs((y_train - y_pred_train)/y_train)*100})
        test_array = pd.DataFrame({'y_test': np.array(y_test),
                                   'y_pred_test': y_pred_test,
                                   'AE': np.abs(y_test - y_pred_test),
                                   'APE': np.abs((y_test - y_pred_test)/y_test)*100})
        
        train_array.to_excel('RS{} {} train.xlsx'.format(rs, model))
        test_array.to_excel('RS{} {} test.xlsx'.format(rs, model))       
       
        # Build DataFrame from evaluation metrics of both training and test datasets
        train_eval = pd.Series({'RMSE': round(np.sqrt(mean_squared_error(y_train, y_pred_train)), 2),
                                'R^2': round(r2_score(y_train, y_pred_train), 2),
                                'MAE': round(sum(train_array['AE'])/len(y_train), 2),
                                'MAPE':round(sum(train_array['APE'])/len(y_train), 2)})
        test_eval = pd.Series({'RMSE': round(np.sqrt(mean_squared_error(y_test, y_pred_test)), 2),
                                'R^2': round(r2_score(y_test, y_pred_test), 2),
                                'MAE': round(sum(test_array['AE'])/len(y_test), 2),
                                'MAPE':round(sum(test_array['APE'])/len(y_test), 2)})
        df2 = pd.DataFrame({'Train': train_eval, 
                            'Test': test_eval})
        # Draw figure for visualiation and comparison
        plt.figure(figsize = (8, 8))
        plt.scatter(y_train, y_pred_train, c='Tab:purple', alpha = 0.6, edgecolors = 'none',
                    label= "Train (RMSE:{} R^2:{} MAE:{} MAPE:{})".format(df2.iloc[0, 0], df2.iloc[1, 0], df2.iloc[2, 0], df2.iloc[3,0]))
        plt.scatter(y_test, y_pred_test, c='Tab:blue', alpha = 0.6, edgecolors = 'none',
                    label= "Test (RMSE:{} R^2:{} MAE:{} MAPE:{})".format(df2.iloc[0, 1], df2.iloc[1, 1], df2.iloc[2, 1], df2.iloc[3, 1]))
        plt.plot([0, 20], [0, 20], c ='orange', ls = '--')
        plt.axis([0, 20, 0, 20])
        plt.xlabel('y_test', fontsize = 18)
        plt.ylabel('y_pred', fontsize = 18)
        plt.title('{}_RS{}'.format(model, rs), fontsize = 22)
        plt.legend(loc="lower right", frameon = False)
        plt.savefig('RS{} {}.png'.format(rs, model), dpi = 300)
        plt.show()
        # Save all results into a word file
        dc = Document()
        dc.add_heading('RS{} {}'.format(str(rs), model), 1)
        dc.add_paragraph('param_grid_{}: {}'.format(model, grid_param))
        dc.add_paragraph('grid_{}.best_params_: {}'.format(model, best_params))
        dc.add_paragraph('grid_{}.best_estimator_: {}'.format(model, best_estimator))
        table = dc.add_table(rows = 1, cols = 3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metrics'
        hdr_cells[1].text = 'Train'
        hdr_cells[2].text = 'Test'
        row_cells = table.add_row().cells
        row_cells[0].text = 'RMSE'
        row_cells[1].text = str(df2.iloc[0, 0])
        row_cells[2].text = str(df2.iloc[0, 1])     
        row_cells = table.add_row().cells
        row_cells[0].text = 'R^2'
        row_cells[1].text = str(df2.iloc[1, 0])
        row_cells[2].text = str(df2.iloc[1, 1])
        row_cells[0].text = 'MAE'
        row_cells[1].text = str(df2.iloc[2, 0])
        row_cells[2].text = str(df2.iloc[2, 1])
        row_cells[0].text = 'MAPE'
        row_cells[1].text = str(df2.iloc[3, 0])
        row_cells[2].text = str(df2.iloc[3, 1])
     
        dc.save('RS{} {}.docx'.format(str(rs), model))            
        return df2

    print('============================================================')
    print(' RS' + str(rs))
    
    #'''
    #M1 ElasticNet Regressor
    model = 'M1 EN'
    en = ElasticNet()
    grid_param_en = {'alpha': [0.001, 0.01, 0.1, 1],
                     'l1_ratio': np.arange(0, 1.1, 0.1)}
    grid_en = GridSearchCV(en, grid_param_en, 
                           cv = cv, scoring = scoring, 
                           refit = True, n_jobs = -1)
    grid_en.fit(X_train_scaled, y_train)
    
    # Use best parameters to retrain the model
    en.set_params(**grid_en.best_params_)
    en.fit(X_train_scaled, y_train)
    y_pred_train_en = en.predict(X_train_scaled)
    y_pred_test_en = en.predict(X_test_scaled)
    df_eval_en = save_evaluate_draw(y_pred_train_en, y_pred_test_en, 
                                    grid_param_en, grid_en.best_params_, 
                                    grid_en.best_estimator_, model)
    print('============================================================')
    print('M1 ElasticNet Regressor')
    print('================================')
    print('1.1 Best parameters:', grid_en.best_params_)
    print('================================')
    print('1.2 Evaluation metrics:')
    print(df_eval_en)
    joblib.dump(en, 'RS{} {}.pkl'.format(rs, model))
    #'''
            
    #'''
    #M2 k-Nearest Neighbors Regressor
    model = 'M2 KNN'
    knn = KNeighborsRegressor()
    grid_param_knn = {'n_neighbors': range(3, 11),
                      'weights': ['uniform', 'distance'],
                      'algorithm': ['ball_tree', 'kd_tree']}
    grid_knn = GridSearchCV(knn, grid_param_knn, 
                            cv = cv, scoring = scoring, 
                            refit = True, n_jobs = -1)
    grid_knn.fit(X_train_scaled, y_train)
    knn.set_params(**grid_knn.best_params_)
    knn.fit(X_train_scaled, y_train)
    y_pred_train_knn = knn.predict(X_train_scaled)
    y_pred_test_knn = knn.predict(X_test_scaled)
    df_eval_knn = save_evaluate_draw(y_pred_train_knn, y_pred_test_knn, 
                                     grid_param_knn, grid_knn.best_params_, 
                                     grid_knn.best_estimator_, model)
    print('============================================================')
    print('M2 KNN Regressor')
    print('================================')
    print('2.1 Best parameters:', grid_knn.best_params_)
    print('================================')
    print('2.2 Evaluation metrics:')
    print(df_eval_knn)
    joblib.dump(knn, 'RS{} {}.pkl'.format(rs, model))
    #'''
    
    #'''
    #3 Support Vector Machine Regressor
    model = 'M3 SVM'
    svr = SVR()
    grid_param_svr = {'kernel': ['poly', 'rbf'], 
                      'C': [1, 1.2, 1.5, 2.0, 3.0, 5.0, 10, 100, 1000],
                      'epsilon': [0.01, 0.05, 0.1, 0.5, 1.0],
                      'degree': range(3, 8)}
    grid_svr = GridSearchCV(svr, grid_param_svr, 
                            cv = cv, scoring = scoring, 
                            refit = True, n_jobs = -1)
    grid_svr.fit(X_train_scaled, y_train)
    svr.set_params(**grid_svr.best_params_)
    svr.fit(X_train_scaled, y_train)
    y_pred_train_svr = svr.predict(X_train_scaled)
    y_pred_test_svr = svr.predict(X_test_scaled)
    df_eval_svr = save_evaluate_draw(y_pred_train_svr, y_pred_test_svr, 
                                     grid_param_svr, grid_svr.best_params_, 
                                     grid_svr.best_estimator_, model)
    print('============================================================')
    print('M3 Support Vector Machine Regressor')
    print('================================')
    print('3.1 Best parameters:', grid_svr.best_params_)
    print('================================')
    print('3.2 Evaluation metrics:')
    print(df_eval_svr)      
    joblib.dump(svr, 'RS{} {}.pkl'.format(rs, model))

    #M4 Decision Tree Rregressor
    model = 'M4 DT'
    dt = DecisionTreeRegressor(random_state = rs)
    grid_param_dt = {'max_depth': range(3, 11)}
    grid_dt = GridSearchCV(dt, grid_param_dt, 
                           cv = cv, scoring = scoring, 
                           refit = True, n_jobs = -1)
    grid_dt.fit(X_train_scaled, y_train)
    dt.set_params(**grid_dt.best_params_)
    dt.fit(X_train_scaled, y_train)
    y_pred_train_dt = dt.predict(X_train_scaled)
    y_pred_test_dt = dt.predict(X_test_scaled)
    df_eval_dt = save_evaluate_draw(y_pred_train_dt, y_pred_test_dt, 
                                    grid_param_dt, grid_dt.best_params_, 
                                    grid_dt.best_estimator_, model)
    print('============================================================')
    print('M4 Decision Tree Regressor')
    print('================================')
    print('4.1 Best parameters:', grid_dt.best_params_)
    print('================================')
    print('4.2 Evaluation metrics:')
    print(df_eval_dt)
    joblib.dump(dt, 'RS{} {}.pkl'.format(rs, model))        
    
    #M5 Random Forest Regressor
    model = 'M5 RF'
    rf = RandomForestRegressor(random_state=rs)
    grid_param_rf = {'max_depth': range(3, 11),
                     'min_samples_split': range(2, 8),
                     'n_estimators': [10, 100, 200, 500, 1000]
                    }
    grid_rf = GridSearchCV(rf, grid_param_rf, 
                           cv = cv, scoring = scoring, 
                           refit = True, n_jobs = -1)
    grid_rf.fit(X_train_scaled, y_train)
    rf.set_params(**grid_rf.best_params_)
    rf.fit(X_train_scaled, y_train)
    y_pred_train_rf = rf.predict(X_train_scaled)
    y_pred_test_rf = rf.predict(X_test_scaled)   
    df_eval_rf = save_evaluate_draw(y_pred_train_rf, y_pred_test_rf, 
                                    grid_param_rf, grid_rf.best_params_, 
                                    grid_rf.best_estimator_, model)
    print('============================================================')
    print('M5 Random Forest Regressor')
    print('================================')
    print('5.1 Best parameters:', grid_rf.best_params_)
    print('================================')
    print('5.2 Evaluation metrics:')
    print(df_eval_rf)  
    joblib.dump(rf, 'RS{} {}.pkl'.format(rs, model))
    #'''
      
    #'''
    #6 XGBoost Regressor
    model = 'M6 XGB'
    xgb = XGBRegressor(objective = 'reg:squarederror', random_state = rs)
    grid_param_xgb = {'max_depth': range(3, 11),
                      'learning_rate': [0.001, 0.01, 0.1, 1],
                      'n_estimators': [10, 100, 200, 500, 1000],
                      'subsample': [0.6, 0.7, 0.8, 0.9, 1.0],
                     }
    grid_xgb = GridSearchCV(xgb, grid_param_xgb, 
                            cv = cv, scoring = scoring, 
                            refit = True, n_jobs = -1)
    grid_xgb.fit(X_train_scaled, y_train)
    xgb.set_params(**grid_xgb.best_params_)
    xgb.fit(X_train_scaled, y_train)
    y_pred_train_xgb = xgb.predict(X_train_scaled)
    y_pred_test_xgb = xgb.predict(X_test_scaled)
    df_eval_xgb = save_evaluate_draw(y_pred_train_xgb, y_pred_test_xgb, 
                                     grid_param_xgb, grid_xgb.best_params_, 
                                     grid_xgb.best_estimator_, model)
    
    # Draw feature imortance of XGBoost
    ax = plot_importance(xgb, importance_type = 'weight')
    fig = ax.figure
    fig.set_size_inches(20, 20)
    fig.savefig('RS{} {} Feat_imp.png'.format(rs, model), dpi = 300)
    plt.show()
    
    # Draw tree
    plot_tree(xgb, num_trees = -1)
    fig = plt.gcf()
    fig.set_size_inches(15, 10)
    plt.savefig('RS{} {} Tree.png'.format(rs, model))
    plt.show()
    
    print('============================================================')
    print('M6 XGBoost Regressor')
    print('================================')
    print('6.1 Best parameters:', grid_xgb.best_params_)
    print('================================')
    print('6.2 Evaluation metrics:')
    print(df_eval_xgb) 
    joblib.dump(xgb, 'RS{} {}.pkl'.format(rs, model))
    #'''
    
    #M7 Lasso
    model = 'M7 Lasso'
    lasso = Lasso()
    grid_param_lasso = {'alpha': np.arange(0, 1, 0.1)}
    grid_lasso = GridSearchCV(lasso, grid_param_lasso, 
                              cv = cv, scoring = scoring, 
                              refit = True, n_jobs = -1)
    grid_lasso.fit(X_train_scaled, y_train)
    lasso.set_params(**grid_lasso.best_params_)
    lasso.fit(X_train_scaled, y_train)
    y_pred_train_lasso = lasso.predict(X_train_scaled)
    y_pred_test_lasso = lasso.predict(X_test_scaled)
    df_eval_lasso = save_evaluate_draw(y_pred_train_lasso, y_pred_test_lasso, 
                                    grid_param_lasso, grid_lasso.best_params_, 
                                    grid_lasso.best_estimator_, model)
    print('============================================================')
    print('M7 Lasso Regressor')
    print('================================')
    print('7.1 Best parameters:', grid_lasso.best_params_)
    print('================================')
    print('7.2 Evaluation metrics:')
    print(df_eval_lasso)
    joblib.dump(lasso, 'RS{} {}.pkl'.format(rs, model))
    #'''
    
    #'''
    #8 Multilayer Perceptron
    model = 'M8 MLP'
    def build_mlp(units = 20, optimizer = 'rmsprop'):
        mlp = Sequential()
        mlp.add(Dense(units = units, activation = 'relu', input_dim = X.shape[1]))
        mlp.add(Dense(units = units, activation = 'relu'))
        mlp.add(Dense(units = 10))
        mlp.add(Dense(units = 1))
        mlp.compile(optimizer = optimizer, 
                    loss='mean_squared_error', 
                    metrics = ['mae'])
        return mlp
    mlp = KerasRegressor(build_fn = build_mlp, verbose = 0)
    grid_param_mlp = {'units': [20, 40, 60],
                      'batch_size': [16, 32, 64],
                      'epochs': [10, 100, 500, 1000],
                      'optimizer': ['rmsprop', 'adam']}
    grid_mlp = GridSearchCV(mlp, grid_param_mlp) 
    grid_mlp.fit(X_train_scaled, y_train)
    y_pred_train_mlp = grid_mlp.predict(X_train_scaled)
    y_pred_test_mlp = grid_mlp.predict(X_test_scaled)
    df_eval_mlp = save_evaluate_draw(y_pred_train_mlp, y_pred_test_mlp, 
                                    grid_param_mlp, grid_mlp.best_params_, 
                                    grid_mlp.best_estimator_, model)
    print('============================================================')
    print('M8 Multilayer Perceptron Regressor')
    print('================================')
    print('8.1 Best parameters:', grid_mlp.best_params_)
    print('================================')
    print('8.2 Evaluation metrics:')
    print(df_eval_mlp)  
    #'''
    

    
    
    
    
    
    
    
    
    
    
    
    
    
    